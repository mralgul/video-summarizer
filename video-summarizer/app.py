from flask import Flask, render_template, request, jsonify, make_response
from PyPDF2 import PdfReader
from pytube import YouTube, extract
import google.generativeai as genai
import os
import re
from io import BytesIO
import docx
from dotenv import load_dotenv
from fpdf import FPDF
import time
import logging
from werkzeug.utils import secure_filename
from concurrent.futures import ThreadPoolExecutor

# Yapılandırma
load_dotenv()
app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB limit
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Güncellenmiş Gemini API Konfigürasyonu
genai.configure(
    api_key=os.getenv("API_KEY"),
    transport='rest',  # v1beta endpoint'i için rest transport kullan
    client_options={
        'api_endpoint': 'https://generativelanguage.googleapis.com/v1beta'
    }
)

# Model tanımı
generation_config = {
    "temperature": 0.3,
    "top_p": 0.95,
    "top_k": 40,
    "max_output_tokens": 2048,
}

safety_settings = [
    {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
]

model = genai.GenerativeModel(
    'gemini-pro',
    generation_config=generation_config,
    safety_settings=safety_settings
)

executor = ThreadPoolExecutor(max_workers=4)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() == 'pdf'

def is_valid_youtube_url(url):
    youtube_regex = (
        r'(https?://)?(www\.)?'
        '(youtube|youtu|youtube-nocookie)\.(com|be)/'
        '(watch\?v=|embed/|v/|.+\?v=)?([^&=%\?]{11})')
    return re.match(youtube_regex, url) is not None

def clean_text(text, max_length=15000):
    if not text:
        return ""
    text = re.sub(r'\s+', ' ', text).strip()
    text = re.sub(r'\[.*?\]', '', text)  # Köşeli parantez içindekileri kaldır
    text = re.sub(r'\(.*?\)', '', text)  # Parantez içindekileri kaldır
    return text[:max_length]

def extract_text_from_pdf(pdf_file):
    try:
        with BytesIO(pdf_file.read()) as f:
            reader = PdfReader(f)
            text = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
            return " ".join(text)
    except Exception as e:
        logger.error(f"PDF Okuma Hatası: {str(e)}")
        return None

def get_youtube_transcript(video_url):
    try:
        yt = YouTube(video_url)
        caption = (yt.captions.get_by_language_code('tr') or 
                  yt.captions.get_by_language_code('en') or 
                  yt.captions.get_by_language_code('a.en'))
        return caption.generate_srt_captions() if caption else ""
    except Exception as e:
        logger.warning(f"Transkript alınamadı: {str(e)}")
        return ""

def generate_summary_prompt(content_type, text):
    return f"""Aşağıdaki {content_type} içeriğini detaylı bir şekilde Türkçe olarak özetle:
    
    {text}

    Özetleme Kuralları:
    1. Ana başlıklar halinde (1. 2. 3. şeklinde) ve her başlık altında 3-5 madde
    2. Her madde maksimum 2 cümle
    3. Teknik terimleri basit ifadelerle açıkla
    4. Önemli noktaları vurgula
    5. Toplam 300-500 kelime arasında olsun
    6. Özeti markdown formatında hazırla (başlıklar için ##, maddeler için -)
    7. Eğitim odaklı bir dil kullan"""

def markdown_to_html(text):
    # Basit markdown'dan HTML'e dönüşüm
    text = re.sub(r'## (.*)', r'<h5 class="mt-4 mb-2 fw-bold">\1</h5>', text)
    text = re.sub(r'- (.*)', r'<div class="mb-2 ps-3"><i class="bi bi-arrow-right me-2"></i>\1</div>', text)
    return text

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/summarize', methods=['POST'])
def summarize_video():
    start_time = time.time()
    try:
        data = request.json
        video_url = data.get('url', '').strip()

        if not video_url:
            return jsonify({'success': False, 'error': 'Lütfen bir YouTube URL girin'}), 400

        if not is_valid_youtube_url(video_url):
            return jsonify({'success': False, 'error': 'Geçersiz YouTube URL formatı'}), 400

        # YouTube bilgilerini al
        yt = YouTube(video_url)
        video_title = clean_text(yt.title) or "Başlıksız Video"
        video_description = clean_text(yt.description) or ""
        
        # Transkripti thread ile paralel al
        transcript_future = executor.submit(get_youtube_transcript, video_url)
        
        # Video bilgilerini hazırla
        full_content = f"BAŞLIK: {video_title}\nAÇIKLAMA: {video_description}"
        
        # Transkripti bekle ve ekle
        transcript = transcript_future.result()
        if transcript:
            full_content += f"\nTRANSCRIPT:\n{clean_text(transcript)}"
        
        # Özetleme işlemi
        response = model.generate_content(generate_summary_prompt("YouTube videosu", full_content))
        
        # Sonuçları formatla
        summary_html = markdown_to_html(response.text)
        
        return jsonify({
            'success': True,
            'title': video_title,
            'summary': summary_html,
            'raw_summary': response.text,
            'thumbnail': yt.thumbnail_url,
            'duration': f"{time.time() - start_time:.2f} saniye"
        })

    except Exception as e:
        logger.error(f"Video Özetleme Hatası: {str(e)}")
        return jsonify({
            'success': False,
            'error': f"Video özetlenirken hata oluştu: {str(e)}"
        }), 500

@app.route('/summarize_pdf', methods=['POST'])
def summarize_pdf():
    start_time = time.time()
    try:
        if 'pdf' not in request.files:
            return jsonify({'success': False, 'error': 'Dosya seçilmedi'}), 400
        
        pdf_file = request.files['pdf']
        if not (pdf_file and allowed_file(pdf_file.filename)):
            return jsonify({'success': False, 'error': 'Sadece PDF dosyaları kabul edilir'}), 400

        # PDF'den metin çıkar
        text = extract_text_from_pdf(pdf_file)
        if not text:
            return jsonify({'success': False, 'error': 'PDF metin içermiyor veya okunamadı'}), 400

        # Özetleme işlemi
        response = model.generate_content(generate_summary_prompt("PDF dokümanı", clean_text(text)))
        
        # Sonuçları formatla
        summary_html = markdown_to_html(response.text)
        
        return jsonify({
            'success': True,
            'title': secure_filename(pdf_file.filename.rsplit('.', 1)[0]),
            'summary': summary_html,
            'raw_summary': response.text,
            'duration': f"{time.time() - start_time:.2f} saniye"
        })

    except Exception as e:
        logger.error(f"PDF Özetleme Hatası: {str(e)}")
        return jsonify({
            'success': False,
            'error': f"PDF özetlenirken hata oluştu: {str(e)}"
        }), 500

@app.route('/download_pdf', methods=['POST'])
def download_pdf():
    try:
        data = request.json
        title = secure_filename(data.get('title', 'Özet')[:100])
        content = data.get('content', '')

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        # Başlık
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(200, 10, txt=title, ln=1, align='C')
        pdf.ln(10)
        
        # İçerik
        pdf.set_font("Arial", size=12)
        for line in content.split('\n'):
            if line.strip():
                pdf.multi_cell(0, 10, txt=line)
                pdf.ln(5)
        
        response = make_response(pdf.output(dest='S').encode('latin1'))
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = f'attachment; filename="{title}_özet.pdf"'
        return response

    except Exception as e:
        logger.error(f"PDF Oluşturma Hatası: {str(e)}")
        return jsonify({
            'success': False,
            'error': f"PDF oluşturulurken hata: {str(e)}"
        }), 500

@app.route('/download_docx', methods=['POST'])
def download_docx():
    try:
        data = request.json
        title = secure_filename(data.get('title', 'Özet')[:100])
        content = data.get('content', '')

        doc = docx.Document()
        doc.add_heading(title, level=1)
        
        for line in content.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        response = make_response(file_stream.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename="{title}_özet.docx"'
        return response

    except Exception as e:
        logger.error(f"DOCX Oluşturma Hatası: {str(e)}")
        return jsonify({
            'success': False,
            'error': f"Word belgesi oluşturulurken hata: {str(e)}"
        }), 500

if __name__ == '__main__':
    try:
        from waitress import serve
        logger.info("Sunucu başlatılıyor...")
        serve(app, host="0.0.0.0", port=5000)
    except ImportError:
        app.run(debug=True)
